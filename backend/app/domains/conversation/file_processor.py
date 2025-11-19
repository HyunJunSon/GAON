import os
import uuid
from typing import BinaryIO, Tuple, List, Dict, Any
from google.cloud import storage
from pypdf import PdfReader
from docx import Document
import io


class FileProcessor:
    def __init__(self, bucket_name: str = "gaon-cloud-data"):
        self.bucket_name = bucket_name
        self.base_path = "user-upload-conv-data"  # 기본 경로
        self.client = storage.Client()
        self.bucket = self.client.bucket(bucket_name)

    def validate_file_content(self, file_content: bytes, file_type: str) -> bool:
        """파일 내용 유효성 검사"""
        try:
            if file_type == "pdf":
                PdfReader(io.BytesIO(file_content))
            elif file_type == "docx":
                Document(io.BytesIO(file_content))
            elif file_type == "txt":
                file_content.decode('utf-8')
            return True
        except Exception:
            return False

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """파일에서 텍스트 추출"""
        try:
            if file_type == "txt":
                return file_content.decode('utf-8')
            
            elif file_type == "pdf":
                reader = PdfReader(io.BytesIO(file_content))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            
            elif file_type == "docx":
                doc = Document(io.BytesIO(file_content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip()
            
            return ""
        except Exception as e:
            raise ValueError(f"텍스트 추출 실패: {str(e)}")

    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """텍스트를 청크로 분할"""
        if not text:
            return []
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # 문장 경계에서 자르기 (마지막 마침표, 느낌표, 물음표 찾기)
            if end < len(text):
                last_sentence_end = max(
                    chunk_text.rfind('.'),
                    chunk_text.rfind('!'),
                    chunk_text.rfind('?'),
                    chunk_text.rfind('\n')
                )
                if last_sentence_end > chunk_size // 2:  # 청크의 절반 이상에서 발견된 경우만
                    chunk_text = chunk_text[:last_sentence_end + 1]
                    end = start + last_sentence_end + 1
            
            chunks.append({
                "chunk_id": chunk_id,
                "text": chunk_text.strip(),
                "start_pos": start,
                "end_pos": end,
                "length": len(chunk_text.strip())
            })
            
            chunk_id += 1
            start = end - overlap  # 오버랩 적용
            
            if start >= len(text):
                break
        
        return chunks

    def upload_to_gcs(self, file_content: bytes, user_id: int, original_filename: str) -> str:
        """GCS에 파일 업로드 (속도 최적화 적용)"""
        try:
            # 고유한 파일 경로 생성
            file_id = str(uuid.uuid4())
            file_extension = original_filename.split('.')[-1].lower()
            # 경로: user-upload-conv-data/conversations/user_123/uuid.txt
            gcs_path = f"{self.base_path}/conversations/user_{user_id}/{file_id}.{file_extension}"
            
            # GCS에 업로드 (폴더 구조는 자동으로 생성됨)
            blob = self.bucket.blob(gcs_path)
            
            # 파일 확장자에 따른 Content-Type 설정
            content_type_map = {
                'webm': 'audio/webm',
                'wav': 'audio/wav',
                'mp3': 'audio/mpeg',
                'm4a': 'audio/mp4',
                'txt': 'text/plain',
                'pdf': 'application/pdf',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            content_type = content_type_map.get(file_extension, 'application/octet-stream')
            
            # 🚀 속도 최적화 설정
            # 1. 청크 크기 최적화 (256KB - 8MB 권장, 기본값보다 큰 값)
            blob.chunk_size = 1024 * 1024 * 2  # 2MB chunks (기본값 대비 향상)
            
            # 2. 큰 파일의 경우 resumable upload 활성화 (자동)
            # 3. 압축 가능한 텍스트 파일의 경우 gzip 압축
            if file_extension in ['txt', 'json', 'csv']:
                import gzip
                compressed_content = gzip.compress(file_content)
                blob.content_encoding = 'gzip'
                blob.upload_from_string(compressed_content, content_type=content_type)
            else:
                blob.upload_from_string(file_content, content_type=content_type)
            
            return gcs_path
        except Exception as e:
            raise ValueError(f"GCS 업로드 실패: {str(e)}")

    def delete_from_gcs(self, gcs_path: str) -> bool:
        """GCS에서 파일 삭제"""
        try:
            blob = self.bucket.blob(gcs_path)
            blob.delete()
            return True
        except Exception:
            return False