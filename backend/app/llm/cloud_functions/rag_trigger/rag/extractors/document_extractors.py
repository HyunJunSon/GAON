"""
문서 형식별 추출 유틸리티 클래스들
"""
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from abc import ABC, abstractmethod

# 로깅 모듈 가져오기
from ..logger import rag_logger

logger = rag_logger


class BaseExtractor(ABC):
    """
    문서 추출기를 위한 추상 베이스 클래스
    """
    
    @abstractmethod
    def extract(self, source: str) -> str:
        """
        문서에서 텍스트를 추출합니다.
        
        Args:
            source: 추출할 소스 (파일 경로 등)
            
        Returns:
            추출된 텍스트
        """
        pass


class PDFExtractor(BaseExtractor):
    """
    PDF 파일에서 텍스트를 추출하는 클래스
    """
    
    def extract(self, source: str) -> str:
        """
        PDF 파일에서 텍스트를 추출합니다.
        
        Args:
            source: PDF 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            import pypdf
            
            with open(source, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                
            logger.info(f"PDF 파일에서 텍스트 추출 완료: {source}")
            return content
        except ImportError:
            raise ImportError("PDF 파일 처리를 위해 'pypdf' 패키지를 설치해야 합니다: pip install pypdf")
        except Exception as e:
            logger.error(f"PDF 파일 추출 오류 {source}: {str(e)}")
            raise


class EPubExtractor(BaseExtractor):
    """
    EPUB 파일에서 텍스트를 추출하는 클래스
    """
    
    def extract(self, source: str) -> str:
        """
        EPUB 파일에서 텍스트를 추출합니다.
        
        Args:
            source: EPUB 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(source)
            
            content = ""
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    content += soup.get_text() + "\n"
            
            logger.info(f"EPUB 파일에서 텍스트 추출 완료: {source}")
            return content
        except ImportError:
            raise ImportError("EPUB 파일 처리를 위해 'ebooklib' 및 'beautifulsoup4' 패키지를 설치해야 합니다: "
                            "pip install ebooklib beautifulsoup4")
        except Exception as e:
            logger.error(f"EPUB 파일 추출 오류 {source}: {str(e)}")
            raise


class TXTExtractor(BaseExtractor):
    """
    일반 텍스트 파일에서 텍스트를 추출하는 클래스
    """
    
    def extract(self, source: str) -> str:
        """
        일반 텍스트 파일에서 텍스트를 추출합니다.
        
        Args:
            source: 텍스트 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            with open(source, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.info(f"텍스트 파일에서 텍스트 추출 완료: {source}")
            return content
        except Exception as e:
            logger.error(f"텍스트 파일 추출 오류 {source}: {str(e)}")
            raise


class MDExtractor(BaseExtractor):
    """
    마크다운 파일에서 텍스트를 추출하는 클래스
    """
    
    def extract(self, source: str) -> str:
        """
        마크다운 파일에서 텍스트를 추출합니다.
        
        Args:
            source: 마크다운 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            with open(source, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            html_content = markdown.markdown(md_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            content = soup.get_text()
            
            logger.info(f"마크다운 파일에서 텍스트 추출 완료: {source}")
            return content
        except ImportError:
            # 마크다운 변환 없이 원본 텍스트 반환
            with open(source, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.warning(f"markdown 또는 beautifulsoup4가 설치되지 않아 마크다운 변환 없이 파일을 로드했습니다: {source}")
            logger.info(f"마크다운 파일에서 텍스트 추출 완료: {source}")
            return content
        except Exception as e:
            logger.error(f"마크다운 파일 추출 오류 {source}: {str(e)}")
            raise


class DOCXExtractor(BaseExtractor):
    """
    DOCX 파일에서 텍스트를 추출하는 클래스
    """
    
    def extract(self, source: str) -> str:
        """
        DOCX 파일에서 텍스트를 추출합니다.
        
        Args:
            source: DOCX 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            from docx import Document
            
            doc = Document(source)
            content = ""
            
            # 단락 텍스트 추출
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # 표의 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + "\n"
            
            logger.info(f"DOCX 파일에서 텍스트 추출 완료: {source}")
            return content
        except ImportError:
            raise ImportError("DOCX 파일 처리를 위해 'python-docx' 패키지를 설치해야 합니다: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX 파일 추출 오류 {source}: {str(e)}")
            raise


class ExtractorFactory:
    """
    파일 형식에 따라 적절한 추출기를 반환하는 팩토리 클래스
    """
    
    def __init__(self):
        self._extractors = {
            '.pdf': PDFExtractor(),
            '.epub': EPubExtractor(),
            '.txt': TXTExtractor(),
            '.md': MDExtractor(),
            '.markdown': MDExtractor(),
            '.docx': DOCXExtractor(),
        }
    
    def get_extractor(self, file_path: str) -> BaseExtractor:
        """
        파일 경로에 따라 적절한 추출기를 반환합니다.
        
        Args:
            file_path: 파일 경로
            
        Returns:
            BaseExtractor 인스턴스
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in self._extractors:
            return self._extractors[file_ext]
        else:
            # 기본 텍스트 추출기 사용
            logger.warning(f"지원되지 않는 파일 형식입니다: {file_ext}. 기본 텍스트 추출기를 사용합니다.")
            return self._extractors['.txt']
    
    def extract_text(self, file_path: str) -> str:
        """
        파일에서 텍스트를 추출합니다.
        
        Args:
            file_path: 추출할 파일 경로
            
        Returns:
            추출된 텍스트
        """
        extractor = self.get_extractor(file_path)
        return extractor.extract(file_path)